"""Generate professionally formatted INTERACT HRM 2.0 Access Control plan DOCX."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(r"C:\Users\Waqas Rafique\Downloads\INTERACT_HRM_2_Professional.docx")

NAVY = RGBColor(0x1A, 0x36, 0x5D)
TEAL = RGBColor(0x00, 0xB8, 0xA9)
BLUE = RGBColor(0x00, 0x52, 0xCC)
GRAY = RGBColor(0x71, 0x80, 0x96)
DARK = RGBColor(0x2D, 0x37, 0x48)


def set_cell_shading(cell, hex_color: str):
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def style_doc(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for level, size, color in [
        ("Heading 1", 18, NAVY),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 12, TEAL),
    ]:
        h = doc.styles[level]
        h.font.name = "Calibri"
        h.font.bold = True
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(16 if level == "Heading 1" else 12)
        h.paragraph_format.space_after = Pt(8)


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INTERACT HRM 2.0")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Centralized Access Control")
    r2.font.size = Pt(20)
    r2.font.color.rgb = BLUE

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Implementation Plan")
    r3.font.size = Pt(16)
    r3.font.color.rgb = TEAL

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('System Control — One-Page Administration for Roles, Permissions & Hierarchy')
    sr.font.size = Pt(11)
    sr.font.italic = True
    sr.font.color.rgb = GRAY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Version 1.1  |  June 2026  |  Interact Global")
    mr.font.size = Pt(10)
    mr.font.color.rgb = GRAY

    doc.add_page_break()


def add_bullets(doc: Document, items: list[str], level: int = 0):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill="1A365D"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
        set_cell_shading(hdr[i], header_fill)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)


def build():
    doc = Document()
    style_doc(doc)
    add_title_page(doc)

    # Executive summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This document defines the implementation plan for Interact HRM 2.0 Centralized Access Control — "
        "a single administrative interface (System Control) where Super Administrators configure roles, "
        "permissions, organizational hierarchy, and feature availability without recurring backend changes."
    )
    doc.add_paragraph(
        "The existing Roles & Permissions page at /admin/roles-permissions is currently a static prototype. "
        "This plan transforms it into a live control center backed by MySQL, with dynamic menus, page guards, "
        "and API-level enforcement."
    )

    doc.add_heading("Business Objectives", level=2)
    add_numbered(
        doc,
        [
            "Create and edit roles (Manager, Team Lead, HR, and custom roles as the organization grows).",
            "Assign feature access via permission checkboxes per role.",
            "Assign employees to roles from a dropdown — access applies on next login.",
            "Enable or disable features globally from the UI when business rules change.",
            "Scale organization structure (departments, teams, new titles) entirely from the frontend.",
        ],
    )

    doc.add_page_break()

    # Section 1
    doc.add_heading("1. Design Principles", level=1)
    doc.add_heading("1.1 Frontend-First, Code-Minimal Approach", level=2)
    doc.add_paragraph(
        "Developers register each new feature once in a central registry file. After deployment and sync, "
        "Super Administrators control visibility and access through the System Control UI — no additional "
        "role-specific code changes are required."
    )
    doc.add_paragraph("Core rule: Features are data, not hardcoded role names scattered across the codebase.")
    add_bullets(
        doc,
        [
            "Avoid: if (role === \"HOD\") show monthly-attendance",
            "Prefer: if (hasPermission(\"attendance.monthly.view\")) show link from database",
        ],
    )

    doc.add_heading("1.2 System Control Architecture", level=2)
    add_code_block(
        doc,
        "System Control Center (/admin/roles-permissions)\n"
        "  Tab 1: Roles          — Create roles, clone templates\n"
        "  Tab 2: Permissions    — Role × feature checkmark matrix\n"
        "  Tab 3: Assign Users   — Org hierarchy (dept, manager, team)\n"
        "  Tab 4: Features       — Global enable/disable toggles\n"
        "  Tab 5: Settings       — Workflow defaults\n"
        "        ↓ saves to MySQL\n"
        "  hrm_roles | hrm_features | hrm_role_permissions | hrm_employees.role_id\n"
        "        ↓ read at login\n"
        "  Dynamic menu | Page guards | API permission checks",
    )

    doc.add_heading("1.3 The Three Pillars", level=2)

    doc.add_heading("Pillar A — Feature Catalog", level=3)
    doc.add_paragraph(
        "Every screen, button, and API action is represented as a row in hrm_features. "
        "Examples include attendance.summary.view, leave.approve.manager, payroll.monthly.edit, "
        "and system.control.access. Features are grouped in the UI by module: "
        "Attendance, Leave, Payroll, Onboarding, Shifts, Reports, and System."
    )

    doc.add_heading("Pillar B — Roles", level=3)
    doc.add_paragraph(
        "Roles are stored in hrm_roles and created from the UI. Built-in templates include "
        "Super Admin, CEO, HR, Manager, Team Lead, Officer, Accountant, and Recruiter. "
        "Custom roles (e.g., Night Shift Manager, Regional HOD) can be cloned from templates."
    )
    add_bullets(
        doc,
        [
            "display_name, slug, description",
            "hierarchy_level — for org chart ordering",
            "data_scope — ALL | DEPARTMENT | TEAM | SELF",
            "portal_type — ADMIN_SHELL | EMPLOYEE_SHELL | LEADER_SHELL",
            "is_system — prevents deletion of core roles",
        ],
    )

    doc.add_heading("Pillar C — User & Org Assignment", level=3)
    doc.add_paragraph(
        "The Assign Users tab defines organizational placement — not just role dropdowns. "
        "It establishes Department → Manager, Team Lead → Team members, and Officer reporting chains. "
        "The Permissions tab controls feature access; Assign Users controls who sees which department "
        "or team data."
    )

    doc.add_page_break()

    # Section 2 UI
    doc.add_heading("2. System Control User Interface", level=1)
    doc.add_paragraph(
        'Recommend renaming the sidebar link from "Roles & Permissions" to "System Control". '
        "Route may remain /admin/roles-permissions or migrate to /admin/system-control."
    )

    tabs = [
        ("Roles", "List and create roles; set portal type, data scope, hierarchy level; clone or delete."),
        ("Permissions", "Role selector with accordion modules; checkmarks per feature; Save / Reset to template."),
        ("Assign Users", "Hierarchy: department managers, team lead membership, employee role and reporting."),
        ("Features", "Global toggles (biometric, Tungsten sync, two-step leave, Team Lead module)."),
        ("Settings", "Default role, leave workflow, session timeout, System Control access policy."),
    ]
    add_table(doc, ["Tab", "Purpose"], tabs)

    doc.add_heading("2.1 Permissions Tab (Detail)", level=2)
    doc.add_paragraph(
        "Super Admin selects a role, expands module groups (Attendance, Leave, Payroll, etc.), "
        "and toggles individual permissions. Quick actions include Select all in group, Clear all, "
        "and Copy from another role. Super Admin and CEO roles display all permissions as locked ON."
    )

    doc.add_heading("2.2 Assign Users Tab (Detail)", level=2)
    doc.add_paragraph(
        "This section defines organizational hierarchy, not just system roles:"
    )
    add_numbered(
        doc,
        [
            "Department → Manager — assigns a department head with full department-scoped visibility.",
            "Team Lead → Team members — groups officers; Team Lead has read-only team attendance view.",
            "Officer → Department + Team Lead — places staff in department and reporting chain.",
        ],
    )
    doc.add_paragraph(
        "Sub-views: Hierarchy view (org chart cards), Department & Manager, Team Lead & members, "
        "All employees (role, department, reports-to, team lead columns)."
    )

    doc.add_page_break()

    # Section 3 Database
    doc.add_heading("3. Database Schema", level=1)
    doc.add_heading("3.1 Core Tables", level=2)
    doc.add_paragraph("The following tables form the foundation of centralized access control.")

    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ["hrm_roles", "Role definitions: slug, portal_type, data_scope, hierarchy_level"],
            ["hrm_features", "Feature catalog: keys, routes, API bindings, global enable flag"],
            ["hrm_role_permissions", "Many-to-many: role_id × feature_key × allowed"],
            ["hrm_employees.role_id", "Links each employee to primary role"],
            ["hrm_teams / hrm_team_members", "Team Lead → officer membership (org hierarchy)"],
        ],
    )

    doc.add_heading("3.2 Optional Tables (Recommended)", level=2)
    add_bullets(
        doc,
        [
            "hrm_role_permission_audit — change log for compliance",
            "hrm_employee_roles — multi-role support (Phase 2)",
        ],
    )

    doc.add_heading("3.3 Seed Data", level=2)
    add_table(
        doc,
        ["Slug", "Display Name", "Portal", "Scope", "Level"],
        [
            ["super_admin", "Super Admin", "admin", "all", "1"],
            ["ceo", "CEO", "admin", "all", "2"],
            ["hr", "HR", "admin", "all", "10"],
            ["manager", "Manager", "admin", "department", "20"],
            ["team_lead", "Team Lead", "leader", "team", "30"],
            ["officer", "Officer", "employee", "self", "90"],
            ["accountant", "Accountant", "admin", "all", "15"],
            ["recruiter", "Recruiter", "admin", "department", "25"],
        ],
    )

    doc.add_page_break()

    # Section 4 Runtime
    doc.add_heading("4. Runtime Access Flow", level=1)
    doc.add_heading("4.1 Login & Session", level=2)
    add_numbered(
        doc,
        [
            "User authenticates via /api/employee-login or admin login.",
            "System loads employee record and joins hrm_roles via role_id.",
            "GET /api/access-control/me returns role, permissions[], features_enabled[], and dynamic menu[].",
            "Session stored in httpOnly cookie or server session (not localStorage alone).",
            "User routed by portal_type: admin → layout-dashboard; leader → leader-dashboard; employee → employee-dashboard.",
        ],
    )

    doc.add_heading("4.2 Data Scope Filtering", level=2)
    add_table(
        doc,
        ["Scope", "Query behavior"],
        [
            ["all", "All employees and departments"],
            ["department", "Filtered to employee_jobs.department_id = user's department"],
            ["team", "Filtered to hrm_team_members for assigned Team Lead"],
            ["self", "Filtered to employee_id = current user only"],
        ],
    )

    doc.add_heading("4.3 Enforcement Layers", level=2)
    add_bullets(
        doc,
        [
            "Dynamic sidebar — menu built from /me response",
            "Page guard — useRequirePermission('feature.key') on each admin page",
            "API guard — requirePermission(req, 'feature.key') with 403 on denial",
        ],
    )

    doc.add_page_break()

    # Section 5 Feature registry
    doc.add_heading("5. Feature Registry", level=1)
    doc.add_paragraph(
        "Single source file: lib/access-control/feature-registry.ts. "
        "POST /api/access-control/sync-registry upserts keys into hrm_features without removing "
        "existing permission assignments. New features appear unchecked for all roles except Super Admin."
    )

    doc.add_heading("5.1 Initial Feature Modules", level=2)
    modules = [
        "Dashboard & Main — dashboard.view, admin.home.view",
        "Leave / PTO — leave.list.view, leave.approve.manager, leave.approve.hr, …",
        "Attendance — summary, manage, monthly, deduction export, Tungsten, employee report",
        "Onboarding & People — employees.add, face.enroll, recruitment, …",
        "Shifts — scheduler and management view/edit",
        "Payroll & Finance — monthly payroll, commissions, advance, loan",
        "Organization — departments, events, calendar, company policy",
        "Team Lead — team dashboard, team attendance/breaks/prayer (view only)",
        "System Control — roles, permissions, features, user assignment, audit",
        "Employee Portal — dashboard, time, leave, credentials (Officer shell)",
    ]
    add_bullets(doc, modules)

    doc.add_heading("5.2 Default Role Templates", level=2)
    add_table(
        doc,
        ["Role", "Suggested permissions"],
        [
            ["Super Admin / CEO", "All features (system-enforced, not editable)"],
            ["HR", "Full HR ops; optional exclusion of system.roles.delete"],
            ["Manager", "Dept-scoped attendance view/edit, leave approve (manager), employee list"],
            ["Team Lead", "Portal self + team.* view permissions"],
            ["Officer", "portal.* self-service only"],
            ["Accountant", "Payroll modules + attendance.monthly.view"],
            ["Recruiter", "recruitment.*, employees.add, employees.list.view"],
        ],
    )

    doc.add_page_break()

    # Section 6 API
    doc.add_heading("6. API Endpoints", level=1)
    doc.add_paragraph("All endpoints under /api/access-control/:")
    api_rows = [
        ["GET /me", "Current user permissions and menu"],
        ["GET/POST/PUT/DELETE /roles", "Role CRUD"],
        ["GET/PUT /roles/:id/permissions", "Permission matrix for role"],
        ["GET /features", "Grouped feature list"],
        ["PUT /features/:key/toggle", "Global feature enable/disable"],
        ["GET/PUT /assignments", "Employee role and org assignments"],
        ["POST /sync-registry", "Sync feature-registry.ts → database"],
        ["GET /audit", "Permission change log"],
    ]
    add_table(doc, ["Endpoint", "Description"], api_rows)

    # Section 7 Migration
    doc.add_heading("7. Migration from Current System", level=1)
    add_numbered(
        doc,
        [
            "Create tables and seed roles/features.",
            "Map hrm_employees.role strings → role_id (BOD/CEO→ceo, HOD→manager, Leader→team_lead, Officer→officer).",
            "Replace static Roles & Permissions page with System Control UI.",
            "Add GET /api/access-control/me on login.",
            "Switch layout-dashboard to dynamic menu (feature flag: USE_DYNAMIC_MENU).",
            "Add requirePermission to critical APIs; roll out to all APIs over 2–3 sprints.",
            "Link or retire hardcoded admin@interact.com credentials.",
        ],
    )

    # Section 8 Phases
    doc.add_heading("8. Implementation Phases", level=1)
    add_table(
        doc,
        ["Phase", "Duration", "Deliverables"],
        [
            ["Phase 1 — Foundation", "1–2 weeks", "SQL, registry, APIs, System Control UI (Roles, Permissions, Assign)"],
            ["Phase 2 — Dynamic UI", "1 week", "Dynamic sidebar, page guards, portal routing, feature toggles"],
            ["Phase 3 — API enforcement", "2 weeks", "requirePermission on all APIs, data scope helpers, audit log"],
            ["Phase 4 — Polish", "Ongoing", "Clone role, bulk assign, multi-role, settings tab, legacy cleanup"],
        ],
    )

    # Section 9 Responsibilities
    doc.add_heading("9. Developer vs Super Admin Responsibilities", level=1)
    add_table(
        doc,
        ["Task", "Developer (code)", "Super Admin (UI)"],
        [
            ["Build new page", "Yes", "—"],
            ["Register feature key in registry", "Yes (one entry)", "—"],
            ["Sync registry to database", "Deploy / button", "Click Sync features"],
            ["Control who opens a page", "—", "Permission checkmarks"],
            ["Create custom role", "—", "Roles tab"],
            ["Assign employee to role", "—", "Assign Users tab"],
            ["Disable module company-wide", "—", "Features toggle"],
            ["Enable two-step leave", "—", "Settings tab"],
        ],
    )

    # Section 10 Success & Risks
    doc.add_heading("10. Success Criteria", level=1)
    add_numbered(
        doc,
        [
            "Super Admin creates a cloned role in under 30 seconds.",
            "Permission changes apply without deployment.",
            "Employee menu updates on next login after role assignment.",
            "API returns 403 when URL is accessed without permission.",
            "Team Lead sees only team-scoped data, not company-wide records.",
        ],
    )

    doc.add_heading("11. Risks & Mitigations", level=1)
    add_table(
        doc,
        ["Risk", "Mitigation"],
        [
            ["Locking out all admins", "Super Admin is_system; break-glass env override"],
            ["Performance on every request", "Cache permissions in session/JWT; invalidate on change"],
            ["Checkbox overload", "Module grouping; clone role; apply template"],
            ["Frontend-only security", "Mandatory requirePermission on all APIs"],
            ["Legacy role === 'HOD' checks", "Phased migration to permission keys"],
        ],
    )

    doc.add_heading("12. Recommended Next Steps", level=1)
    add_numbered(
        doc,
        [
            "Approve plan and execute Phase 1.",
            "Create scripts/hrm_access_control.sql.",
            "Implement lib/access-control/feature-registry.ts.",
            "Build app/api/access-control/* routes.",
            "Deploy System Control UI (MVP: Roles, Permissions, Assign Users).",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("Related documentation: docs/HRM_Organizational_Hierarchy_Technical_Document.txt")
    for r in p.runs:
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = GRAY

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
